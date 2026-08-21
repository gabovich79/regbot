import os
import re
import hashlib
import fitz  # pymupdf
from docx import Document
import httpx
from bs4 import BeautifulSoup
from config import DOCUMENTS_DIR, ORIGINALS_DIR

os.makedirs(DOCUMENTS_DIR, exist_ok=True)
os.makedirs(ORIGINALS_DIR, exist_ok=True)


from services.token_utils import estimate_tokens


def clean_text(text: str) -> str:
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"[ \t]+", " ", text)
    return text.strip()


def extract_pdf_pages(file_path: str) -> list[dict]:
    """Extract text page-by-page so citations can point to the source page."""
    doc = fitz.open(file_path)
    try:
        return [
            {"page_number": index, "text": clean_text(page.get_text())}
            for index, page in enumerate(doc, start=1)
        ]
    finally:
        doc.close()


def extract_pdf(file_path: str) -> str:
    doc = fitz.open(file_path)
    pages = []
    for page in doc:
        pages.append(page.get_text())
    doc.close()
    return clean_text("\n".join(pages))


def extract_docx(file_path: str) -> str:
    doc = Document(file_path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    return clean_text("\n".join(paragraphs))


def extract_pdf_bytes_pages(content: bytes) -> list[dict]:
    """Extract in-memory PDF content with stable 1-based page numbers."""
    doc = fitz.open(stream=content, filetype="pdf")
    try:
        return [
            {"page_number": index, "text": clean_text(page.get_text())}
            for index, page in enumerate(doc, start=1)
        ]
    finally:
        doc.close()


def extract_pdf_bytes(content: bytes) -> str:
    doc = fitz.open(stream=content, filetype="pdf")
    pages = []
    for page in doc:
        pages.append(page.get_text())
    doc.close()
    return clean_text("\n".join(pages))


def extract_docx_bytes(content: bytes) -> str:
    import io
    doc = Document(io.BytesIO(content))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    return clean_text("\n".join(paragraphs))


def normalize_source_url(source_ref: str) -> str | None:
    """Extract the actual URL from raw UI/export strings stored in legacy rows."""
    match = re.search(r"https?://[^\s`]+", source_ref or "")
    return match.group(0) if match else None


async def fetch_url_document(url: str) -> tuple[str, list[dict] | None, bytes, str]:
    """Fetch a URL and retain page boundaries when the response is a PDF."""
    async with httpx.AsyncClient(follow_redirects=True, timeout=30) as client:
        response = await client.get(url)
        response.raise_for_status()
        content_type = response.headers.get("content-type", "").lower()
        normalized_url = str(response.url).lower()

        if "pdf" in content_type or normalized_url.endswith(".pdf"):
            pages = extract_pdf_bytes_pages(response.content)
            return "\n\n".join(page["text"] for page in pages), pages, response.content, "pdf"
        if "wordprocessingml" in content_type or "docx" in content_type or normalized_url.endswith(".docx"):
            return extract_docx_bytes(response.content), None, response.content, "docx"
        if "msword" in content_type or normalized_url.endswith(".doc"):
            raise ValueError("legacy .doc URL requires manual conversion before page-aware re-ingestion")

        soup = BeautifulSoup(response.text, "html.parser")
        for tag in soup(["script", "style", "nav", "footer", "header"]):
            tag.decompose()
        text = soup.get_text(separator="\n")
        return clean_text(text), None, response.content, "html"


async def fetch_url_text(url: str) -> str:
    text, _, _, _ = await fetch_url_document(url)
    return text


def extract_gdrive_file_id(url: str) -> str | None:
    patterns = [
        r"/file/d/([a-zA-Z0-9_-]+)",
        r"id=([a-zA-Z0-9_-]+)",
        r"/d/([a-zA-Z0-9_-]+)",
    ]
    for pattern in patterns:
        match = re.search(pattern, url)
        if match:
            return match.group(1)
    return None


async def fetch_gdrive_text(url: str) -> str:
    file_id = extract_gdrive_file_id(url)
    if not file_id:
        raise ValueError("לא ניתן לחלץ File ID מהקישור של Google Drive")

    download_url = f"https://drive.google.com/uc?export=download&id={file_id}"
    async with httpx.AsyncClient(follow_redirects=True, timeout=30) as client:
        response = await client.get(download_url)
        response.raise_for_status()
        content_type = response.headers.get("content-type", "")

        if "pdf" in content_type:
            return extract_pdf_bytes(response.content)
        elif "document" in content_type or "docx" in content_type:
            return extract_docx_bytes(response.content)
        else:
            return clean_text(response.text)


def save_original_document(doc_id: int, extension: str, content: bytes) -> tuple[str, str]:
    """Persist the immutable source bytes and return path plus SHA-256 checksum."""
    safe_extension = re.sub(r"[^a-z0-9]", "", extension.lower()) or "bin"
    path = os.path.join(ORIGINALS_DIR, f"{doc_id}.{safe_extension}")
    with open(path, "wb") as file:
        file.write(content)
    return path, hashlib.sha256(content).hexdigest()


def save_document_text(doc_id: int, text: str) -> str:
    path = os.path.join(DOCUMENTS_DIR, f"{doc_id}.txt")
    with open(path, "w", encoding="utf-8") as f:
        f.write(text)
    return path


def load_document_text(text_path: str) -> str:
    with open(text_path, "r", encoding="utf-8") as f:
        return f.read()


def delete_document_file(text_path: str):
    if os.path.exists(text_path):
        os.remove(text_path)
